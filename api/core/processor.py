import re
from typing import Dict, Any
from io import BytesIO
from html import escape, unescape
from html.parser import HTMLParser

from docx import Document
from docx.shared import RGBColor

GMAIL_FONT_STACKS = {
    "Sans Serif": "Arial, Helvetica, sans-serif",
    "Serif": "'Times New Roman', Times, serif",
    "等寬字型": "'Courier New', Courier, monospace",
    "微軟正黑體": "'Microsoft JhengHei', 'PingFang TC', 'Noto Sans TC', sans-serif",
    "新細明體": "'PMingLiU', 'MingLiU', 'Noto Serif TC', serif",
    "細明體": "'MingLiU', 'PMingLiU', 'Noto Serif TC', serif",
    "寬": "'Arial Black', 'Impact', sans-serif",
    "窄": "'Arial Narrow', 'Helvetica Neue Condensed', sans-serif",
    "Comic Sans MS": "'Comic Sans MS', 'Comic Sans', cursive",
    "Garamond": "Garamond, 'Times New Roman', serif",
    "Georgia": "Georgia, 'Times New Roman', serif",
    "Tahoma": "Tahoma, 'Segoe UI', sans-serif",
    "Trebuchet MS": "'Trebuchet MS', 'Segoe UI', sans-serif",
    "Verdana": "Verdana, 'Segoe UI', sans-serif",
}

DEFAULT_TEMPLATE_TEXT = "請在這裡編輯郵件內容"


def resolve_gmail_font(font_key: str | None) -> str:
    if not font_key:
        return GMAIL_FONT_STACKS["Sans Serif"]
    key = font_key.strip()
    return GMAIL_FONT_STACKS.get(key, GMAIL_FONT_STACKS["Sans Serif"])


def create_default_template_html(base_font_family: str | None = None) -> str:
    font_family = base_font_family or resolve_gmail_font(None)
    return (
        f"<div style=\"font-family: {font_family}; line-height: 1.6; color: #333;\">"
        f"<p>{escape(DEFAULT_TEMPLATE_TEXT)}</p>"
        "</div>"
    )


def resolve_template_html(
    *,
    docx_content: bytes | None = None,
    template_html: str | None = None,
    base_font_family: str | None = None,
) -> str:
    if template_html and template_html.strip():
        return template_html
    if docx_content is not None:
        return convert_docx_to_html(docx_content, base_font_family=base_font_family)
    return create_default_template_html(base_font_family=base_font_family)

def convert_docx_to_html(
    file_bytes: bytes,
    log_colors: bool = False,
    log_fonts: bool = False,
    base_font_family: str | None = None,
) -> str:
    """
    將 .docx 轉為適合 Email 的 HTML 格式，保留常見的行內樣式（粗體、斜體、底線、文字顏色）
    以及無序/有序清單。
    log_colors / log_fonts 用來決定是否在 stdout 列出偵測到的色彩與字型，預設不列印。
    """
    doc = Document(BytesIO(file_bytes))
    numbering_root = doc.part.numbering_part.element if doc.part.numbering_part else None

    def log_detected_colors():
        colors = []
        for p_idx, paragraph in enumerate(doc.paragraphs):
            for r_idx, run in enumerate(paragraph.runs):
                col = run.font.color.rgb
                if col:
                    snippet = run.text.replace("\n", " ")[:30]
                    colors.append((p_idx, r_idx, str(col), snippet))
        if colors:
            print("[convert_docx_to_html] detected run colors:")
            for p_idx, r_idx, col, snippet in colors:
                print(f"  p{p_idx} r{r_idx}: #{col} text='{snippet}'")
        else:
            print("[convert_docx_to_html] no run-level colors found")

    if log_colors:
        log_detected_colors()

    def get_font_name(run):
        # 優先讀取 rFonts（常見於 Word 預設「新細明體」等）
        rpr = run._r.rPr
        if rpr is not None and getattr(rpr, "rFonts", None) is not None:
            rfonts = rpr.rFonts
            for attr in ("eastAsia", "ascii", "hAnsi", "cs"):
                val = getattr(rfonts, attr, None)
                if val:
                    return val
        # 其次讀取直接套用的 run.font.name
        return run.font.name

    def log_detected_fonts():
        fonts = []
        for p_idx, paragraph in enumerate(doc.paragraphs):
            for r_idx, run in enumerate(paragraph.runs):
                fname = get_font_name(run)
                if fname:
                    snippet = run.text.replace("\n", " ")[:30]
                    fonts.append((p_idx, r_idx, fname, snippet))
        if fonts:
            print("[convert_docx_to_html] detected run fonts:")
            for p_idx, r_idx, fname, snippet in fonts:
                print(f"  p{p_idx} r{r_idx}: font='{fname}' text='{snippet}'")
        else:
            print("[convert_docx_to_html] no run-level font names found")

    if log_fonts:
        log_detected_fonts()

    def get_list_tag(paragraph) -> str:
        """
        根據 numFmt 決定使用 <ul> 或 <ol>。若判斷不到，預設用 <ul>。
        """
        if numbering_root is None:
            return "ul"
        num_pr = paragraph._p.pPr.numPr
        num_id = num_pr.numId.val
        ilvl = num_pr.ilvl.val if num_pr.ilvl is not None else 0
        num_el = numbering_root.xpath(f".//w:num[@w:numId='{num_id}']")
        if not num_el:
            return "ul"
        abstract_id = num_el[0].xpath("w:abstractNumId/@w:val")[0]
        fmt = numbering_root.xpath(
            f".//w:abstractNum[@w:abstractNumId='{abstract_id}']/w:lvl[@w:ilvl='{ilvl}']/w:numFmt/@w:val"
        )
        if fmt and fmt[0] != "bullet":
            return "ol"
        return "ul"

    def get_run_style(run):
        color = run.font.color.rgb
        return (bool(run.bold), bool(run.italic), bool(run.underline), str(color) if color else None)

    def text_to_html(text_value: str, style: tuple[bool, bool, bool, str | None]) -> str:
        text = escape(text_value)
        if not text:
            return ""
        # 保留連續空白
        text = text.replace("  ", "&nbsp;&nbsp;")

        is_bold, is_italic, is_underline, color = style
        if is_bold:
            text = f"<strong>{text}</strong>"
        if is_italic:
            text = f"<em>{text}</em>"
        if is_underline:
            text = f"<u>{text}</u>"

        if color:
            text = f'<span style="color: #{color}">{text}</span>'
        return text

    html_parts = []
    in_list = False
    current_list_tag = None

    def paragraph_is_list(p) -> bool:
        return p._p.pPr is not None and p._p.pPr.numPr is not None

    for paragraph in doc.paragraphs:
        is_list = paragraph_is_list(paragraph)

        if is_list:
            list_tag = get_list_tag(paragraph)
            if not in_list or list_tag != current_list_tag:
                if in_list:
                    html_parts.append(f"</{current_list_tag}>")
                html_parts.append(f"<{list_tag}>")
                in_list = True
                current_list_tag = list_tag
        else:
            if in_list:
                html_parts.append(f"</{current_list_tag}>")
                in_list = False
                current_list_tag = None

        segments: list[str] = []
        current_style = None
        current_text = ""
        for run in paragraph.runs:
            if not run.text:
                continue
            run_style = get_run_style(run)
            if current_style is None:
                current_style = run_style
                current_text = run.text
                continue
            if run_style == current_style:
                current_text += run.text
                continue
            segments.append(text_to_html(current_text, current_style))
            current_style = run_style
            current_text = run.text
        if current_style is not None:
            segments.append(text_to_html(current_text, current_style))

        content = "".join(segments)
        if not content:
            continue

        if is_list:
            html_parts.append(f"<li>{content}</li>")
        else:
            html_parts.append(f"<p>{content}</p>")

    if in_list and current_list_tag:
        html_parts.append(f"</{current_list_tag}>")

    font_family = base_font_family or resolve_gmail_font(None)
    wrapped_html = (
        f"<div style=\"font-family: {font_family}; line-height: 1.6; color: #333;\">"
        + "".join(html_parts)
        + "</div>"
    )
    return wrapped_html

def inject_variables(html_template: str, row_data: Dict[str, Any]) -> str:
    """
    將 HTML 模板中的 {{變數}} 替換為 Excel 的資料內容
    支援自定義欄位，只要 Excel 表頭名稱與 {{}} 內一致即可。
    """
    # 允許數字欄位（例如 {{1}}）也能對應到 Excel 的欄位名稱
    normalized_row = {str(k).strip(): v for k, v in row_data.items()}
    replaced_keys: list[str] = []

    def normalize_placeholder_key(raw_key: str) -> str:
        return unescape(re.sub(r"<[^>]+>", "", raw_key)).strip()

    # 使用正則表達式尋找 {{Key}} 並從 row_data 抓取對應的 Value
    def replace_match(match):
        key = normalize_placeholder_key(match.group(1))
        value = normalized_row.get(key)
        if value is None:
            return match.group(0)
        replaced_keys.append(key)
        return escape(str(value))

    # 匹配 {{ variable_name }} 格式；若 Word run 被切開，中間可能夾 HTML tag。
    pattern = r"\{\{\s*(.*?)\s*\}\}"
    rendered = re.sub(pattern, replace_match, html_template, flags=re.DOTALL)

    unresolved = sorted(
        {
            normalize_placeholder_key(match)
            for match in re.findall(pattern, rendered, flags=re.DOTALL)
            if normalize_placeholder_key(match)
        }
    )
    if replaced_keys or unresolved:
        print(
            "[inject_variables] replaced:",
            ", ".join(sorted(set(replaced_keys))) or "(none)",
            "| unresolved:",
            ", ".join(unresolved) or "(none)",
            "| available headers:",
            ", ".join(sorted(normalized_row.keys())),
        )

    return rendered


class _DocxHtmlParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.paragraphs: list[list[dict[str, object]]] = [[]]
        self.style_stack: list[dict[str, object]] = [{}]

    def _current_paragraph(self) -> list[dict[str, object]]:
        return self.paragraphs[-1]

    def _current_style(self) -> dict[str, object]:
        return self.style_stack[-1]

    def _push_paragraph(self) -> None:
        if self.paragraphs and not self.paragraphs[-1]:
            return
        self.paragraphs.append([])

    def _push_text(self, text: str) -> None:
        if not text:
            return
        self._current_paragraph().append({"text": text, "style": dict(self._current_style())})

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        attrs_dict = dict(attrs)
        next_style = dict(self._current_style())
        if tag in {"strong", "b"}:
            next_style["bold"] = True
        elif tag in {"em", "i"}:
            next_style["italic"] = True
        elif tag == "u":
            next_style["underline"] = True
        elif tag == "span":
            style_attr = attrs_dict.get("style") or ""
            color_match = re.search(r"color\s*:\s*#?([0-9a-fA-F]{6})", style_attr)
            if color_match:
                next_style["color"] = color_match.group(1)

        if tag in {"p", "div", "li"}:
            self._push_paragraph()
        elif tag == "br":
            self._push_text("\n")

        self.style_stack.append(next_style)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"p", "div", "li"}:
            self._push_paragraph()
        if len(self.style_stack) > 1:
            self.style_stack.pop()

    def handle_data(self, data: str) -> None:
        self._push_text(data)


def export_html_to_docx_bytes(html: str) -> bytes:
    parser = _DocxHtmlParser()
    parser.feed(html)

    document = Document()
    for segments in parser.paragraphs:
        if not segments:
            continue
        paragraph = document.add_paragraph()
        for segment in segments:
            text = str(segment.get("text", ""))
            style = segment.get("style", {})
            pieces = text.split("\n")
            for idx, piece in enumerate(pieces):
                if idx > 0:
                    paragraph.add_run().add_break()
                if not piece:
                    continue
                run = paragraph.add_run(piece)
                if isinstance(style, dict):
                    if style.get("bold"):
                        run.bold = True
                    if style.get("italic"):
                        run.italic = True
                    if style.get("underline"):
                        run.underline = True
                    color = style.get("color")
                    if isinstance(color, str) and re.fullmatch(r"[0-9a-fA-F]{6}", color):
                        run.font.color.rgb = RGBColor.from_string(color.upper())

    if not document.paragraphs:
        document.add_paragraph("")

    output = BytesIO()
    document.save(output)
    return output.getvalue()
